__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
import argparse
import os
import shutil
from langchain.schema.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from get_embedding_function import get_embedding_function
import pdfplumber
from docx import Document as DocxDocument
from tqdm import tqdm
import sys

# Absolute path for Chroma database
CHROMA_PATH = os.path.abspath("chroma")
DATA_PATH = "data"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--reset", action="store_true", help="Reset the database.")
    args = parser.parse_args()

    # Ensure Chroma path exists and is writable
    ensure_chroma_path()

    if args.reset:
        print("Clearing database...")
        clear_database()

    documents = []
    documents += load_pdf_documents()
    documents += load_word_documents()

    if not documents:
        print("No documents found in the data folder.")
        return

    chunks = split_documents(documents)
    db = add_to_chroma(chunks)

    # Quick test query
    test_query(db)


def ensure_chroma_path():
    try:
        os.makedirs(CHROMA_PATH, exist_ok=True)
        test_file = os.path.join(CHROMA_PATH, "test.db")
        with open(test_file, "w") as f:
            f.write("ok")
        os.remove(test_file)
    except Exception as e:
        print(f"ERROR: Chroma path '{CHROMA_PATH}' is not writable: {e}")
        sys.exit(1)


def load_pdf_documents():
    documents = []
    for file in os.listdir(DATA_PATH):
        if file.lower().endswith(".pdf"):
            pdf_path = os.path.join(DATA_PATH, file)
            with pdfplumber.open(pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        documents.append(Document(
                            page_content=text,
                            metadata={"source": file, "page": i}
                        ))
    print(f"Loaded {len(documents)} PDF pages as documents.")
    return documents


def load_word_documents():
    documents = []
    for file in os.listdir(DATA_PATH):
        if file.lower().endswith(".docx") or file.lower().endswith(".doc"):
            doc_path = os.path.join(DATA_PATH, file)
            try:
                doc = DocxDocument(doc_path)
                for i, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text:
                        documents.append(Document(
                            page_content=text,
                            metadata={"source": file, "page": i}
                        ))
            except Exception as e:
                print(f"Failed to load {file}: {e}")
    print(f"Loaded {len(documents)} Word paragraphs as documents.")
    return documents


def split_documents(documents: list[Document]):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=80,
        length_function=len,
        is_separator_regex=False,
    )
    chunks = text_splitter.split_documents(documents)
    print(f"Split documents into {len(chunks)} chunks.")
    return chunks


def add_to_chroma(chunks: list[Document]):
    try:
        db = Chroma(
            persist_directory=CHROMA_PATH,
            embedding_function=get_embedding_function()
        )
    except Exception as e:
        print(f"Failed to initialize ChromaDB: {e}")
        print("Attempting to clear corrupted database and retry...")
        clear_database()
        db = Chroma(
            persist_directory=CHROMA_PATH,
            embedding_function=get_embedding_function()
        )

    chunks_with_ids = calculate_chunk_ids(chunks)

    existing_items = db.get(include=[])
    existing_ids = set(existing_items["ids"])
    print(f"Existing document count in DB: {len(existing_ids)}")

    new_chunks = [chunk for chunk in chunks_with_ids if chunk.metadata["id"] not in existing_ids]

    if new_chunks:
        print(f"Adding {len(new_chunks)} new chunks to the database...")
        new_chunk_ids = [chunk.metadata["id"] for chunk in new_chunks]
        db.add_documents(new_chunks, ids=new_chunk_ids)  # Batch insert
        print("Finished updating the database.")
    else:
        print("No new documents to add.")

    return db


def calculate_chunk_ids(chunks):
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source", "unknown")
        page = chunk.metadata.get("page", 0)
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk.metadata["id"] = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id

    return chunks


def clear_database():
    if os.path.exists(CHROMA_PATH):
        shutil.rmtree(CHROMA_PATH)
        os.makedirs(CHROMA_PATH, exist_ok=True)
        print("Database cleared.")



if __name__ == "__main__":
    main()

